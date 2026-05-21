"""
Multi-format document loading + cleaning pipeline for the FPViber RAG engine.

Supports .txt, .pdf and .docx files. Designed to handle technical drone
documentation that mixes prose, tables and mathematical formulas.

Cleaning pipeline strategy
--------------------------
1. Extract raw text per file (format-specific):
   - PDF  : pypdf, page-by-page, tables flattened with `|` separators.
   - DOCX : python-docx, walks paragraphs and tables in document order.
   - TXT  : direct read.
2. Normalise whitespace, fix soft hyphens that PDFs love to insert.
3. Strip page numbers and repeated headers / footers detected across pages.
4. Drop short noise lines (single letters, ornamental dashes, etc.).
5. Preserve tables (lines containing the `|` column separator) and math
   blocks (LaTeX-style `$...$`, `\\(...\\)`, `\\[...\\]`) intact - the
   chunker downstream is informed not to split inside them.
"""

from __future__ import annotations

import os
import re
from collections import Counter
from dataclasses import dataclass


SUPPORTED_EXTENSIONS = (".txt", ".pdf", ".docx")


# ----------------------------------------------------------------------
# Data classes
# ----------------------------------------------------------------------

@dataclass
class LoadedDocument:
    """One source document after cleaning, before chunking."""
    source: str          # filename (e.g. "manual.pdf")
    text: str            # cleaned full text


# ----------------------------------------------------------------------
# Per-format raw extractors
# ----------------------------------------------------------------------

def _extract_txt(path: str) -> list[str]:
    with open(path, "r", encoding="utf-8", errors="ignore") as fh:
        return [fh.read()]


def _extract_pdf(path: str) -> list[str]:
    """Return a list with one string per page, tables flattened in-line."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "pypdf is required for PDF support. Install with `pip install pypdf`."
        ) from exc

    reader = PdfReader(path)
    pages: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        pages.append(text)
    return pages


_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _linearize_omml(omath_elem) -> str:
    """Best-effort linearisation of Word OMML into LaTeX-like text."""
    parts: list[str] = []
    for elem in omath_elem.iter():
        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
        if tag == "t" and elem.text:
            parts.append(elem.text.strip())
        elif tag == "f":
            parts.append("/")
        elif tag in ("sup", "sSub", "sSubSup"):
            parts.append("^")
        elif tag in ("sub",):
            parts.append("_")
    formula = " ".join(parts)
    return re.sub(r"\s+", " ", formula).strip()


def _paragraph_with_math(para) -> str:
    """Merge plain runs and inline OMML from one DOCX paragraph."""
    from docx.oxml.ns import qn

    bits: list[str] = []
    for child in para._element.iterchildren():
        if child.tag == qn("w:r"):
            texts = [t.text for t in child.iter(qn("w:t")) if t.text]
            if texts:
                bits.append("".join(texts))
        elif child.tag.endswith("oMath") or child.tag.endswith("oMathPara"):
            formula = _linearize_omml(child)
            if formula:
                bits.append(f"${formula}$")
    if bits:
        return "".join(bits).strip()
    return para.text.strip()


def _extract_docx(path: str) -> list[str]:
    """Walk the docx body in order; preserve tables and OMML formulas."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install with `pip install python-docx`."
        ) from exc

    doc = Document(path)
    body = doc.element.body
    pieces: list[str] = []

    paragraphs = {p._p: p for p in doc.paragraphs}
    tables = {t._tbl: t for t in doc.tables}

    p_tag = qn("w:p")
    tbl_tag = qn("w:tbl")

    for child in body.iterchildren():
        if child.tag == p_tag:
            para = paragraphs.get(child)
            if para is None:
                continue
            block = _paragraph_with_math(para)
            if block:
                pieces.append(block)
            for omath in child.iter(f"{{{_M_NS}}}oMath"):
                formula = _linearize_omml(omath)
                if formula and f"${formula}$" not in block:
                    pieces.append(f"${formula}$")
        elif child.tag == tbl_tag:
            table = tables.get(child)
            if table is None:
                continue
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    pieces.append(row_text)
    return ["\n".join(pieces)]


_EXTRACTORS = {
    ".txt": _extract_txt,
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


# ----------------------------------------------------------------------
# Cleaning
# ----------------------------------------------------------------------

_PAGE_NUM_PATTERNS = [
    re.compile(r"^\s*page\s+\d+(\s*(/|of)\s*\d+)?\s*$", re.IGNORECASE),
    re.compile(r"^\s*-\s*\d+\s*-\s*$"),
    re.compile(r"^\s*\d{1,4}\s*$"),
]

_NOISE_PATTERNS = [
    re.compile(r"^[\W_]{2,}$"),     # only punctuation / ornaments
]

# Standalone equation / calculation lines (PDF often emits these without $ delimiters)
_EQUATION_LINE = re.compile(
    r"^(?:"
    r"\$[^$]+\$|"
    r"\\\[[\s\S]+\\\]|"
    r"(?:\(\d+\)\s*)?[A-Za-z][\w]*\s*=\s*[^=].{2,}|"
    r".{0,40}(?:thrust|torque|power|efficiency|TWR|KV|RPM|current|voltage)"
    r"[^=]{0,30}=\s*[\d\w+\-*/^().%]+"
    r")$",
    re.IGNORECASE,
)

_MATH_SYMBOLS = re.compile(r"[≤≥±×÷√∑∫∞≈≠°^_\d]")

_EQUATION_CONT = re.compile(
    r"^(?:[\d\w+\-*/^().%≤≥±×÷√∑∫∞≈≠°\s]+|"
    r"(?:where|and|with|for)\b.+)$",
    re.IGNORECASE,
)


def _fix_hyphenation(text: str) -> str:
    """Join 'multi-\nline' words that were broken by PDF line wrapping."""
    return re.sub(r"-\n(\w)", r"\1", text)


def _normalise_whitespace(text: str) -> str:
    text = text.replace("\u00ad", "")           # soft hyphen
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _is_page_number(line: str) -> bool:
    return any(p.match(line) for p in _PAGE_NUM_PATTERNS)


def _is_noise(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False  # empty lines are kept; they delimit paragraphs
    if len(stripped) <= 2 and not stripped.isalnum():
        return True
    return any(p.match(stripped) for p in _NOISE_PATTERNS)


def _detect_running_headers_footers(pages: list[str], min_repeat: int = 3) -> set[str]:
    """
    Find lines that repeat at the top or bottom of many pages - these are
    almost certainly headers/footers and should be stripped.
    """
    if len(pages) < min_repeat:
        return set()

    candidates: Counter[str] = Counter()
    for page in pages:
        lines = [ln.strip() for ln in page.splitlines() if ln.strip()]
        for ln in lines[:2] + lines[-2:]:
            if 3 <= len(ln) <= 120:
                candidates[ln] += 1

    threshold = max(min_repeat, int(len(pages) * 0.5))
    return {line for line, n in candidates.items() if n >= threshold}


def _clean_lines(text: str, banned: set[str]) -> str:
    cleaned_lines: list[str] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip() in banned:
            continue
        if _is_page_number(line):
            continue
        if _is_noise(line):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)


def _looks_like_equation(line: str) -> bool:
    s = line.strip()
    if len(s) < 4:
        return False
    if s.startswith("$") or s.startswith("\\["):
        return True
    if _EQUATION_LINE.match(s):
        return True
    if "=" in s and (_MATH_SYMBOLS.search(s) or re.search(r"\b[A-Za-z]\w*\s*=", s)):
        return len(s) < 200
    return False


def _preserve_math_formulas(text: str) -> str:
    """
    Wrap detected equation lines in \\[ ... \\] so the chunker treats them as
    atomic math blocks. Merge PDF line-breaks that split formulas across rows.
    """
    lines = text.split("\n")
    out: list[str] = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            out.append(lines[i])
            i += 1
            continue

        if _looks_like_equation(stripped):
            merged = stripped
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    break
                if _looks_like_equation(nxt) and not _EQUATION_CONT.match(nxt):
                    break
                if _EQUATION_CONT.match(nxt) or (
                    len(nxt) < 80 and ("=" in nxt or _MATH_SYMBOLS.search(nxt))
                ):
                    merged += " " + nxt
                    j += 1
                else:
                    break
            if not (merged.startswith("$") or merged.startswith("\\[")):
                merged = f"\\[{merged}\\]"
            out.append(merged)
            i = j
            continue

        out.append(lines[i])
        i += 1

    return "\n".join(out)


def clean_text(pages: list[str]) -> str:
    """Run the full cleaning pipeline across the raw page list."""
    fixed_pages = [_fix_hyphenation(_normalise_whitespace(p)) for p in pages]
    banned = _detect_running_headers_footers(fixed_pages)
    joined = "\n\n".join(fixed_pages)
    cleaned = _clean_lines(joined, banned).strip()
    return _preserve_math_formulas(cleaned)


# ----------------------------------------------------------------------
# Public loader
# ----------------------------------------------------------------------

def load_documents(folder: str) -> list[LoadedDocument]:
    """
    Load every supported file in `folder` and return a list of
    cleaned `LoadedDocument` objects. Raises if the folder is missing
    or contains no supported documents.
    """
    if not os.path.exists(folder):
        raise FileNotFoundError(
            f"Folder '{folder}' does not exist. Create it and add documents "
            f"(.txt / .pdf / .docx) inside."
        )

    documents: list[LoadedDocument] = []
    for file_name in sorted(os.listdir(folder)):
        ext = os.path.splitext(file_name)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue

        path = os.path.join(folder, file_name)
        try:
            raw_pages = _EXTRACTORS[ext](path)
        except ImportError:
            raise
        except Exception as exc:
            print(f"[loader] WARN: failed to read {file_name}: {exc}", flush=True)
            continue

        text = clean_text(raw_pages)
        if not text:
            print(f"[loader] WARN: {file_name} is empty after cleaning", flush=True)
            continue

        documents.append(LoadedDocument(source=file_name, text=text))

    if not documents:
        raise ValueError(
            f"No usable documents found in '{folder}'. "
            f"Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    return documents
